import os
from flask import Flask, render_template, request, send_file, jsonify
from werkzeug.utils import secure_filename
from PIL import Image
import pandas as pd
from pdf2image import convert_from_path
from pdf2docx import Converter
from docx import Document
from docx2pdf import convert as docx_to_pdf
from fpdf import FPDF
import PyPDF2

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['CONVERTED_FOLDER'] = 'converted'

# Allowed input formats
ALLOWED_EXTENSIONS = {'pdf', 'csv', 'xlsx', 'txt', 'png', 'jpg', 'jpeg', 'docx'}

# Valid conversions map
VALID_CONVERSIONS = {
    'pdf':   ['txt', 'png', 'jpg', 'docx'],
    'csv':   ['pdf', 'xlsx', 'txt'],
    'xlsx':  ['pdf', 'csv', 'txt'],
    'txt':   ['pdf', 'docx'],
    'png':   ['jpg', 'pdf'],
    'jpg':   ['png', 'pdf'],
    'jpeg':  ['png', 'pdf'],
    'docx':  ['pdf', 'txt']
}

# ----------------- Utility Functions -----------------
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def pdf_to_txt(input_path, output_path):
    from PyPDF2 import PdfReader
    text = ""
    reader = PdfReader(input_path)
    for page in reader.pages:
        extracted = page.extract_text()
        if extracted:
            text += extracted + "\n"
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text)

def pdf_to_docx(input_path, output_path):
    cv = Converter(input_path)
    cv.convert(output_path, start=0, end=None)
    cv.close()

def txt_to_docx(input_path, output_path):
    with open(input_path, "r", encoding="utf-8") as f:
        text = f.read()
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    doc.save(output_path)

def docx_to_txt(input_path, output_path):
    doc = Document(input_path)
    full_text = "\n".join([p.text for p in doc.paragraphs])
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(full_text)

def txt_to_pdf(input_path, output_path):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Times", size=12)

    with open(input_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    for line in lines:
        stripped = line.strip()
        if not stripped:
            pdf.ln(6)
            continue
        if stripped.isupper() and len(stripped) <= 60:
            pdf.set_font("Times", style="B", size=14)
            pdf.multi_cell(0, 8, stripped, align="C")
            pdf.ln(4)
            pdf.set_font("Times", size=12)
        else:
            pdf.multi_cell(0, 6, stripped)

    pdf.output(output_path)

def images_to_pdf(input_path, output_path):
    img = Image.open(input_path).convert("RGB")
    img.save(output_path, "PDF", resolution=100.0)

# ----------------- Conversion Handler -----------------
def convert_file(input_path, output_path, input_ext, output_ext):
    # PDF conversions
    if input_ext == 'pdf' and output_ext == 'txt':
        pdf_to_txt(input_path, output_path)
    elif input_ext == 'pdf' and output_ext in ['png', 'jpg']:
        images = convert_from_path(input_path)
        if images:
            images[0].save(output_path, output_ext.upper())
    elif input_ext == 'pdf' and output_ext == 'docx':
        pdf_to_docx(input_path, output_path)

    # DOCX conversions
    elif input_ext == 'docx' and output_ext == 'txt':
        docx_to_txt(input_path, output_path)
    elif input_ext == 'docx' and output_ext == 'pdf':
        temp_dir = os.path.dirname(output_path)
        os.makedirs(temp_dir, exist_ok=True)
        docx_to_pdf(input_path, temp_dir)
        generated_pdf = os.path.join(temp_dir, os.path.splitext(os.path.basename(input_path))[0] + ".pdf")
        os.rename(generated_pdf, output_path)

    # TXT conversions
    elif input_ext == 'txt' and output_ext == 'docx':
        txt_to_docx(input_path, output_path)
    elif input_ext == 'txt' and output_ext == 'pdf':
        txt_to_pdf(input_path, output_path)

    # CSV/XLSX conversions
    elif input_ext in ['csv', 'xlsx'] and output_ext in ['csv', 'xlsx', 'pdf', 'txt']:
        if input_ext == 'csv':
            df = pd.read_csv(input_path, encoding="utf-8")
        elif input_ext == 'xlsx':
            df = pd.read_excel(input_path)
        if output_ext == 'csv':
            df.to_csv(output_path, index=False, encoding="utf-8")
        elif output_ext == 'xlsx':
            df.to_excel(output_path, index=False)
        elif output_ext == 'txt':
            df.to_csv(output_path, sep="\t", index=False, encoding="utf-8")
        elif output_ext == 'pdf':
            html = df.to_html(index=False)
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=8)
            for line in html.splitlines():
                pdf.multi_cell(0, 5, line)
            pdf.output(output_path)

    # Image conversions
    elif input_ext in ['png', 'jpg', 'jpeg'] and output_ext in ['png', 'jpg']:
        img = Image.open(input_path)
        img.save(output_path)
    elif input_ext in ['png', 'jpg', 'jpeg'] and output_ext == 'pdf':
        images_to_pdf(input_path, output_path)

    else:
        raise ValueError("Unsupported conversion requested.")

# ----------------- Routes -----------------
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/get_formats/<ext>')
def get_formats(ext):
    ext = ext.lower()
    formats = VALID_CONVERSIONS.get(ext, [])
    return jsonify(formats)

@app.route('/', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return "No file part", 400
    file = request.files['file']
    if file.filename == '':
        return "No selected file", 400
    if not allowed_file(file.filename):
        return "Invalid file type", 400

    target_format = request.form.get('format')
    if not target_format:
        return "No target format selected", 400

    filename = secure_filename(file.filename)
    input_ext = filename.rsplit('.', 1)[1].lower()

    if target_format not in VALID_CONVERSIONS.get(input_ext, []):
        return "Invalid conversion type", 400

    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    os.makedirs(app.config['CONVERTED_FOLDER'], exist_ok=True)

    input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(input_path)

    base_name = os.path.splitext(filename)[0]
    output_filename = f"{base_name}.{target_format}"
    output_path = os.path.join(app.config['CONVERTED_FOLDER'], output_filename)

    try:
        convert_file(input_path, output_path, input_ext, target_format)
    except Exception as e:
        return f"Error during conversion: {e}", 500

    return send_file(output_path, as_attachment=True)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)